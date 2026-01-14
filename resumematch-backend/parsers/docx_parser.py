from docx import Document
from typing import Dict, List, Any
import io
import logging

logger = logging.getLogger(__name__)

class DocxParser:
    def __init__(self):
        pass
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extract text from a DOCX file
        """
        try:
            doc = Document(file_path)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            return "\n".join(paragraphs).strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise
    
    @staticmethod
    def extract_text_from_bytes(docx_bytes: bytes) -> str:
        """
        Extract text from DOCX bytes
        """
        try:
            doc = Document(io.BytesIO(docx_bytes))
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            return "\n".join(paragraphs).strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX bytes: {str(e)}")
            raise
    
    @staticmethod
    def extract_tables(file_path: str) -> List[List[Dict[str, str]]]:
        """
        Extract tables from a DOCX file
        """
        try:
            doc = Document(file_path)
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = {}
                    for i, cell in enumerate(row.cells):
                        row_data[f"column_{i}"] = cell.text.strip()
                    table_data.append(row_data)
                tables_data.append(table_data)
            return tables_data
        except Exception as e:
            logger.error(f"Error extracting tables from DOCX: {str(e)}")
            return []
    
    @staticmethod
    def extract_document_properties(file_path: str) -> Dict[str, Any]:
        """
        Extract document properties from DOCX file
        """
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            return {
                "title": core_props.title or "",
                "subject": core_props.subject or "",
                "author": core_props.author or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "last_modified_by": core_props.last_modified_by or "",
                "revision": core_props.revision or 0,
            }
        except Exception as e:
            logger.error(f"Error extracting DOCX properties: {str(e)}")
            return {}